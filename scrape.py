import os
import time
import re
import requests
from bs4 import BeautifulSoup
from bs4 import Comment
from urllib.parse import urljoin, urlparse, urlunparse
from io import BytesIO
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
import fitz  # PyMuPDF
from docx import Document
import csv

# === 設定 ===
ENABLE_OCR = False  # 画像OCRをONにする場合 True, OFFにする場合 False
ENABLE_PDF = False   # TrueならPDFもダウンロードしてテキスト抽出

# 除外したいディレクトリパス
exclude_paths = [
    '/en', '/pressroom',  # 例）/en/ 配下や pressroom 配下は除外
]

# 対象にしたいディレクトリパス（空リストなら全体対象）
include_only_prefix = []

# OCR用 言語設定（日本語＋英語）
ocr_lang = 'jpn+eng'

# --- クリーンアップ＆補助関数 ---
def sanitize_text(text):
    return re.sub(r'[\\/*?:"<>|]', "_", text)

def normalize_url(url):
    parsed = urlparse(url)
    path = parsed.path.rstrip('/')
    normalized = urlunparse((parsed.scheme, parsed.netloc, path, '', '', ''))
    return normalized

def clean_text(text):
    if not text:
        return ''
    text = text.replace('\x00', '')
    text = text.encode('utf-8', 'ignore').decode('utf-8', 'ignore')

    def is_valid_xml_char(c):
        codepoint = ord(c)
        return (
            codepoint == 0x9 or
            codepoint == 0xA or
            codepoint == 0xD or
            (0x20 <= codepoint <= 0xD7FF) or
            (0xE000 <= codepoint <= 0xFFFD) or
            (0x10000 <= codepoint <= 0x10FFFF)
        )

    text = ''.join(c for c in text if is_valid_xml_char(c))
    return text

def should_visit(url):
    path = urlparse(url).path
    for exclude in exclude_paths:
        if path.startswith(exclude):
            return False
    if include_only_prefix:
        return any(path.startswith(prefix) for prefix in include_only_prefix)
    return True

def scrape_website(start_url, output_file, exclude_paths=None, enable_ocr=False, enable_pdf=False, include_only_prefix=None, progress_callback=None):
    if exclude_paths is None:
        exclude_paths = []
    if include_only_prefix is None:
        include_only_prefix = []
    visited = set()
    to_visit = [start_url]
    domain = urlparse(start_url).netloc
    doc = Document()
    is_first_page = True
    processed_pages = 0

    def should_visit(url):
        path = urlparse(url).path
        for exclude in exclude_paths:
            if path.startswith(exclude):
                return False
        if include_only_prefix:
            return any(path.startswith(prefix) for prefix in include_only_prefix)
        return True

    while to_visit:
        url = to_visit.pop()
        normalized_current_url = normalize_url(url)
        if normalized_current_url in visited:
            continue
        visited.add(normalized_current_url)

        try:
            response = requests.get(url)
            response.raise_for_status()
            response.encoding = response.apparent_encoding
        except Exception as e:
            print(f"エラー: {url} - {e}")
            continue

        soup = BeautifulSoup(response.text, 'html.parser')

        # --- <header>タグとその関連要素を完全に除去 ---
        # headerタグ自体を除去
        for header_tag in soup.find_all('header'):
            header_tag.decompose()
        
        # headerに関連するクラスやIDを持つ要素も除去
        header_selectors = [
            {'class_': ['header', 'site-header', 'main-header', 'global-header']},
            {'id': ['header', 'site-header', 'main-header', 'global-header']},
            {'role': 'banner'},
            {'class_': ['header-nav', 'header-menu', 'header-wrapper']},
            {'id': ['header-nav', 'header-menu', 'header-wrapper']}
        ]
        
        for selector in header_selectors:
            for tag in soup.find_all(attrs=selector):
                tag.decompose()

        # --- グローバルナビゲーション除去 ---
        for nav in soup.find_all('nav'):
            nav.decompose()
        # <header>タグ内のリストやナビゲーション要素は既に除去される

        unwanted_selectors = [
            {'id': 'global-nav'},
            {'id': 'gnav'},
            {'class': 'global-nav'},
            {'class': 'header-nav'},
            {'class': 'menu'},
            {'class': 'nav'},
            {'id': 'header'},
            {'id': 'footer'},
        ]
        for selector in unwanted_selectors:
            for tag in soup.find_all(attrs=selector):
                tag.decompose()

        # ★ ディスクリプションだけ先に取得
        description_text = None
        description_tag = soup.find('meta', attrs={'name': 'description'})
        if description_tag and description_tag.get('content'):
            description_text = description_tag.get('content').strip()

        # ★ コメントアウトを削除
        for comment in soup.find_all(string=lambda text: isinstance(text, Comment)):
            comment.extract()

        # ★ 不要なmeta, script, style, noscript, iframe, link, svgタグを削除
        for tag in soup(['meta', 'script', 'style', 'noscript', 'iframe', 'link', 'svg']):
            tag.decompose()

        # タイトル取得
        title_tag = soup.find('title')
        title = title_tag.get_text(strip=True) if title_tag else "no_title"
        clean_title = sanitize_text(title)

        # ページ区切り（最初の1ページ目はスキップ）
        if not is_first_page:
            doc.add_page_break()
        else:
            is_first_page = False
        doc.add_heading(clean_text(clean_title), level=1)

        # ディスクリプションがあれば出力
        if description_text:
            doc.add_paragraph(clean_text(f"【ディスクリプション】\n{description_text}"))

        # 本文エリア
        main_content = soup.find('main')
        if main_content:
            content_source = main_content
        else:
            content_source = soup

        # 本文のすべてのテキスト抽出
        for element in content_source.descendants:
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                text = element.get_text(strip=True)
                if text:
                    level = int(element.name[1])
                    doc.add_heading(clean_text(text), level=level)
            elif element.name == 'li':
                text = element.get_text(strip=True)
                if text:
                    doc.add_paragraph(clean_text(text), style='List Bullet')
            elif element.name is None:
                text = element.strip()
                if text:
                    doc.add_paragraph(clean_text(text))

        # 画像からOCR抽出 (ON/OFF)
        if enable_ocr and OCR_AVAILABLE:
            ocr_texts = []
            for img_tag in soup.find_all('img'):
                img_url = img_tag.get('src')
                if img_url:
                    img_full_url = urljoin(url, img_url)
                    if img_full_url.lower().endswith('.svg'):
                        continue
                    try:
                        img_response = requests.get(img_full_url)
                        img_response.raise_for_status()
                        img = Image.open(BytesIO(img_response.content))
                        text_from_image = pytesseract.image_to_string(img, lang=ocr_lang).strip()
                        if text_from_image:
                            ocr_texts.append(text_from_image)
                    except Exception as e:
                        print(f"画像エラー: {img_full_url} - {e}")
                        continue
            if ocr_texts:
                doc.add_heading("画像から抽出されたテキスト", level=2)
                for ocr_text in ocr_texts:
                    doc.add_paragraph(clean_text(ocr_text))
        elif enable_ocr and not OCR_AVAILABLE:
            doc.add_paragraph("※ OCR機能は現在の環境では利用できません")

        # PDFからテキスト抽出 (ON/OFF)
        if enable_pdf:
            for link_tag in soup.find_all('a', href=True):
                href = link_tag['href']
                if href.lower().endswith('.pdf'):
                    full_url = urljoin(url, href)
                    try:
                        pdf_response = requests.get(full_url)
                        pdf_response.raise_for_status()
                        pdf_bytes = BytesIO(pdf_response.content)
                        doc_pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
                        extracted_text = ""
                        for page in doc_pdf:
                            extracted_text += page.get_text()
                        if extracted_text.strip():
                            doc.add_heading("PDFから抽出されたテキスト", level=2)
                            doc.add_paragraph(clean_text(extracted_text.strip()))
                    except Exception as e:
                        print(f"PDFエラー: {full_url} - {e}")
                        continue

        # === 🆕 リンク収集 ===
        for link_tag in soup.find_all('a', href=True):
            href = link_tag['href']
            full_url = urljoin(url, href)
            normalized_url = normalize_url(full_url)
            if (
                urlparse(normalized_url).netloc == domain and
                normalized_url not in visited and
                should_visit(normalized_url) and
                not normalized_url.lower().endswith('.pdf')
            ):
                to_visit.append(normalized_url)

        # ★ ページごとに進捗を表示・報告
        processed_pages += 1
        current_total = processed_pages + len(to_visit)
        if progress_callback:
            progress_callback(done=processed_pages, total=current_total)
        
        path = urlparse(url).path
        print(f"書き込み完了: {path} - {clean_title}")
        time.sleep(1)

    doc.save(output_file)
    print(f"全ページを {output_file} にまとめました！")
    return output_file

def list_all_urls(start_url, exclude_paths=None, include_only_prefix=None, progress_callback=None):
    if exclude_paths is None:
        exclude_paths = []
    if include_only_prefix is None:
        include_only_prefix = []
    visited = set()
    to_visit = [start_url]
    domain = urlparse(start_url).netloc
    url_list = []
    total = 0

    def should_visit(url):
        path = urlparse(url).path
        for exclude in exclude_paths:
            if path.startswith(exclude):
                return False
        if include_only_prefix:
            return any(path.startswith(prefix) for prefix in include_only_prefix)
        return True

    while to_visit:
        url = to_visit.pop()
        normalized_current_url = normalize_url(url)
        if normalized_current_url in visited:
            continue
        visited.add(normalized_current_url)
        url_list.append(normalized_current_url)
        total += 1
        if progress_callback:
            progress_callback(done=total, total=len(visited) + len(to_visit))
        try:
            response = requests.get(url)
            response.raise_for_status()
        except Exception as e:
            print(f"エラー: {url} - {e}")
            continue
        soup = BeautifulSoup(response.text, 'html.parser')
        for link_tag in soup.find_all('a', href=True):
            href = link_tag['href']
            full_url = urljoin(url, href)
            normalized_url = normalize_url(full_url)
            if (
                urlparse(normalized_url).netloc == domain and
                normalized_url not in visited and
                should_visit(normalized_url) and
                not normalized_url.lower().endswith('.pdf')
            ):
                to_visit.append(normalized_url)
    return url_list

def list_all_urls_with_stats(start_url, output_file, exclude_paths=None, include_only_prefix=None, progress_callback=None):
    if exclude_paths is None:
        exclude_paths = []
    if include_only_prefix is None:
        include_only_prefix = []
    visited = set()
    to_visit = [start_url]
    domain = urlparse(start_url).netloc
    url_rows = []  # 各ページの情報
    dir_stats = {}  # ディレクトリごとの統計 {'/service/': {'page': 0, 'pdf': 0}, ...}
    total_pages = 0
    total_pdfs = 0

    def should_visit(url):
        path = urlparse(url).path
        for exclude in exclude_paths:
            if path.startswith(exclude):
                return False
        if include_only_prefix:
            return any(path.startswith(prefix) for prefix in include_only_prefix)
        return True

    def get_directory(path):
        # 例: /service/abc → /service/
        if not path or path == '/':
            return '/'
        parts = path.strip('/').split('/')
        return '/' + parts[0] + '/'

    while to_visit:
        url = to_visit.pop()
        normalized_current_url = normalize_url(url)
        if normalized_current_url in visited:
            continue
        visited.add(normalized_current_url)
        path = urlparse(normalized_current_url).path
        directory = get_directory(path)
        is_pdf = 1 if normalized_current_url.lower().endswith('.pdf') else 0
        url_rows.append([normalized_current_url, directory, is_pdf])
        # ディレクトリ統計
        if directory not in dir_stats:
            dir_stats[directory] = {'page': 0, 'pdf': 0}
        dir_stats[directory]['page'] += 1
        if is_pdf:
            dir_stats[directory]['pdf'] += 1
            total_pdfs += 1
        total_pages += 1
        # 進捗コールバック（処理済み件数と残り件数の合計で計算）
        if progress_callback:
            current_total = total_pages + len(to_visit)
            progress_callback(done=total_pages, total=current_total)
        try:
            response = requests.get(url)
            response.raise_for_status()
        except Exception as e:
            print(f"エラー: {url} - {e}")
            continue
        soup = BeautifulSoup(response.text, 'html.parser')
        for link_tag in soup.find_all('a', href=True):
            href = link_tag['href']
            full_url = urljoin(url, href)
            normalized_url = normalize_url(full_url)
            if (
                urlparse(normalized_url).netloc == domain and
                normalized_url not in visited and
                should_visit(normalized_url)
            ):
                to_visit.append(normalized_url)
    # CSV出力
    with open(output_file, 'w', encoding='utf-8', newline='') as f:
        writer = csv.writer(f)
        # ページ一覧
        writer.writerow(['url', 'directory', 'is_pdf'])
        for row in url_rows:
            writer.writerow(row)
        # 空行
        writer.writerow([])
        # ディレクトリ統計
        writer.writerow(['directory', 'page_count', 'pdf_count'])
        for directory, stats in dir_stats.items():
            writer.writerow([directory, stats['page'], stats['pdf']])
        # 空行
        writer.writerow([])
        # サイト全体統計
        writer.writerow(['total_pages', 'total_pdfs'])
        writer.writerow([total_pages, total_pdfs])
    return output_file